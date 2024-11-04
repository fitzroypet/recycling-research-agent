from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import logging
import re

class DocumentFormatter:
    def __init__(self, location, recycling_data=None):
        self.doc = Document()
        self.location = location
        self.recycling_data = recycling_data
        self._setup_styles()
    
    def _add_section_break(self):
        """Add a section break (paragraph) to the document"""
        self.doc.add_paragraph()
    
    def _setup_styles(self):
        """Setup custom styles for the document"""
        styles = {
            'Custom Heading 1': {'size': 16, 'bold': True, 'color': RGBColor(0, 102, 0)},
            'Custom Normal': {'size': 11},
            'Custom Bullet': {'size': 11}
        }
        
        for name, props in styles.items():
            style = self.doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.size = Pt(props['size'])
            if props.get('bold'):
                style.font.bold = True
            if props.get('color'):
                style.font.color.rgb = props['color']
    
    def format_report(self, content):
        # Add title
        title = self.doc.add_heading(f'Recycling Services Analysis: {self.location}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self._add_section_break()
        
        # Extract sections using regex to handle both markdown and plain text
        sections = re.split(r'\*\*|##\s+', content)
        sections = [s.strip() for s in sections if s.strip()]
        
        for section in sections:
            # Skip empty sections
            if not section:
                continue
            
            # Check if this is a heading
            lines = section.split('\n', 1)
            if len(lines) == 2:
                heading, content = lines
                
                # Add heading
                section_heading = self.doc.add_heading(heading.strip(), level=1)
                section_heading.style = self.doc.styles['Custom Heading 1']
                
                # Process content
                content = content.strip()
                if content:
                    # Handle bullet points
                    for line in content.split('\n'):
                        line = line.strip()
                        if line.startswith('-') or line.startswith('•'):
                            p = self.doc.add_paragraph(
                                line[1:].strip(),
                                style='Custom Bullet'
                            )
                            p.style.paragraph_format.left_indent = Inches(0.25)
                        elif line:  # Skip empty lines
                            self.doc.add_paragraph(line, style='Custom Normal')
                
                self._add_section_break()
        
        return self.doc
    
    def save(self, filename):
        try:
            self.doc.save(filename)
            logging.info(f"Document saved successfully: {filename}")
            return True
        except Exception as e:
            logging.error(f"Error saving document: {str(e)}")
            return False